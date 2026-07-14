import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Task 9: Hyperparameter Tuning Submission', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Objective
doc.add_heading('Objective', level=1)
doc.add_paragraph('Systematically tune hyperparameters to squeeze validated performance out of the model without overfitting, and confirm the gains on a held-out test set.')

# GitHub Repository
doc.add_heading('GitHub Repository Access', level=1)
p = doc.add_paragraph()
p.add_run('IMPORTANT: ').bold = True
p.add_run('The repository is confirmed to be set to PUBLIC. The mentor will be able to access the code and verify the implementation evidence without issues.')

doc.add_heading('Relevant Files for Task 9 Verification:', level=2)
doc.add_paragraph('src/train_model.py: Contains the RandomizedSearchCV pipeline and automatic metric generation.', style='List Bullet')
doc.add_paragraph('metrics.json: The automatically generated output file containing the final metrics.', style='List Bullet')
doc.add_paragraph('experiment_log.md: The historical log updated with the Task 9 run.', style='List Bullet')
doc.add_paragraph('models/rf_job_match_pipeline_tuned.pkl: The final tuned model artifact.', style='List Bullet')

# Methodology & Execution
doc.add_heading('Methodology & Execution', level=1)
doc.add_paragraph('Hyperparameters Chosen: We targeted the most impactful parameters for a Random Forest: n_estimators (number of trees), max_depth (tree depth), and min_samples_split (node splitting rules) to balance complexity and prevent overfitting.', style='List Number')
doc.add_paragraph('Search Strategy & CV Scheme: We utilized RandomizedSearchCV with 15 parameter combinations across a 3-fold cross-validation (cv=3) scheme.', style='List Number')
doc.add_paragraph('Scoring Metric: We evaluated the models using roc_auc, ensuring the model is robust at separating classes across various thresholds.', style='List Number')
doc.add_paragraph('Validation & Confirmation: The optimal configuration achieved a Cross-Validation ROC-AUC of 0.8393. We then confirmed the performance on our strict held-out test set (75 samples), achieving a Test ROC-AUC of 0.8523 and an Accuracy of 77.33%.', style='List Number')
doc.add_paragraph('Artifact Discipline: Directly addressing the Task 8 mentor feedback, src/train_model.py was explicitly modified to automatically generate and save metrics.json at the end of the main() execution. We have committed both this script and the generated metrics to the repository.', style='List Number')

# Best Config
doc.add_heading('Best Configuration Discovered', level=1)
doc.add_paragraph('n_estimators: 50', style='List Bullet')
doc.add_paragraph('min_samples_split: 10', style='List Bullet')
doc.add_paragraph('max_depth: None', style='List Bullet')

# Pitfalls Avoided
doc.add_heading('Pitfalls Successfully Avoided', level=1)
doc.add_paragraph('Tuning on the test set: The test set was strictly partitioned via train_test_split before tuning began. RandomizedSearchCV only ever saw the training data.', style='List Bullet')
doc.add_paragraph('Reporting CV best as final without test confirmation: The reported final metrics (in metrics.json and the experiment log) are evaluated explicitly on the unseen X_test dataset, confirming that the CV gains held up in reality.', style='List Bullet')
doc.add_paragraph('Searching params that don\'t matter: We kept the search space constrained to structural parameters that actively prevent overfitting, rather than wasting compute on negligible hyperparameter tweaks.', style='List Bullet')

doc.save(r'c:\Users\777Pr\Downloads\Task9_Submission.docx')
print("Successfully generated Word document!")
