import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Task 19: Model Serialization and Versioning Submission', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Objective
doc.add_heading('Objective', level=1)
doc.add_paragraph('Serialize the final trained machine learning pipeline (preprocessing + model) into a production-ready, versioned artifact that can be loaded and used for inference in any application, while adding structured Pydantic input verification.')

# GitHub Repository
doc.add_heading('GitHub Repository Access', level=1)
p = doc.add_paragraph()
p.add_run('IMPORTANT: ').bold = True
p.add_run('The repository is confirmed to be set to PUBLIC. The mentor will be able to access the code and verify the implementation evidence without issues.')

doc.add_heading('Relevant Files for Task 19 Verification:', level=2)
doc.add_paragraph('src/model_serialization.py: Contains the robust ModelLoader class, the predict_match function with Pydantic validation, and the core joblib serialization logic.', style='List Bullet')
doc.add_paragraph('models/job_match_pipeline_v1.0.pkl: The finalized bundled pipeline artifact representing the model and data preprocessor.', style='List Bullet')
doc.add_paragraph('artifacts/metadata.json: Comprehensive snapshot tracking model provenance, environment configurations, and metrics.', style='List Bullet')
doc.add_paragraph('artifacts/model_validation.json: Output proving identical probability mapping before and after deserialization.', style='List Bullet')
doc.add_paragraph('reports/compatibility_report.md: Versioning documentation bounding expected Python and Scikit-Learn targets.', style='List Bullet')

# Methodology & Execution
doc.add_heading('Methodology & Execution', level=1)
doc.add_paragraph('End-to-end Bundling: We utilized joblib to extract the completely assembled XGBoost predictor—including all custom preprocessing steps and scaling boundaries—freezing them together as one unit so preprocessing steps are never skipped during live inference.', style='List Number')
doc.add_paragraph('Secure Inference Guardrails: We built a Pydantic schema (JobMatchInput) directly into the prediction interface ensuring only structured, in-bound datatypes reach the model payload, immediately catching bad inference calls.', style='List Number')
doc.add_paragraph('Reproducibility Proof: A fresh loading environment was spun up using JobMatchModelLoader, comparing predictions against the active memory pipeline. Results matched identically with <1e-6 deviation.', style='List Number')
doc.add_paragraph('Metadata and Environment Locking: A detailed metadata.json file was generated binding the model configuration, dataset history, validation metrics, Git commit hash, and Scikit-Learn versions to ensure deployment traceability.', style='List Number')

doc.save(r'c:\Users\777Pr\Downloads\Task19_Submission.docx')
print("Successfully generated Word document for Task 19!")
