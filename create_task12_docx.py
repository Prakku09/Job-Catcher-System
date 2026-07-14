import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('Task 12: Production-Grade Binary Classification', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Direct Access & Source Code
doc.add_heading('1. Source Code & GitHub Repository Access', level=1)
p = doc.add_paragraph()
p.add_run('The GitHub repository is PUBLIC and fully accessible at: ').bold = True
p.add_run('https://github.com/Prakku09/Job-Catcher-System\n')

p.add_run('\nDirect links to the implementation evidence:\n').bold = True
doc.add_paragraph('Production Pipeline Script: https://github.com/Prakku09/Job-Catcher-System/blob/main/src/production_pipeline.py', style='List Bullet')
doc.add_paragraph('Live Verification Results: https://github.com/Prakku09/Job-Catcher-System/blob/main/reports/live_predictions.csv', style='List Bullet')
doc.add_paragraph('Threshold Optimization (Operating Point): https://github.com/Prakku09/Job-Catcher-System/blob/main/reports/operating_point.json', style='List Bullet')
doc.add_paragraph('Probability Calibration Output: https://github.com/Prakku09/Job-Catcher-System/blob/main/plots/calibration_curve.png', style='List Bullet')

# 2. Dependency Tracking
doc.add_heading('2. Dependency Tracking', level=1)
p = doc.add_paragraph('We strictly track all environment dependencies to ensure full reproducibility in production. The exact versions of Pandas, NumPy, Scikit-Learn, and Joblib used to compile the model are automatically saved in the ')
p.add_run('reports/dependencies.json').bold = True
p.add_run(' artifact. Furthermore, every single generated report is embedded with metadata capturing the Python version, Scikit-Learn version, Git commit hash, and a UTC timestamp.')

# 3. Failure Handling & Edge Cases
doc.add_heading('3. Failure Handling Mechanisms', level=1)
p = doc.add_paragraph('Our production pipeline is designed to gracefully handle failures, as documented thoroughly in ')
p.add_run('docs/failure_handling.md').bold = True
p.add_run(' and ')
p.add_run('docs/edge_cases.md').bold = True
p.add_run(':\n')

doc.add_paragraph('Missing Data / Model Files: The API boundary features startup checks. If production_binary_classifier.pkl is missing, it logs a CRITICAL alert rather than failing silently.', style='List Bullet')
doc.add_paragraph('Invalid Payloads: We utilize strict schema validation at the API layer. Empty payloads ({}), missing predictive features, or negative values (e.g., negative experience years) are rejected with an HTTP 422 Unprocessable Entity before reaching the Scikit-Learn pipeline.', style='List Bullet')
doc.add_paragraph('Unexpected Exceptions: The model prediction call is wrapped in a try-except block to intercept unexpected mathematical faults (like NaN propagation), returning a safe HTTP 500 without crashing the worker.', style='List Bullet')

# 4. Calibration & Operating Point
doc.add_heading('4. Probability Calibration & Optimal Threshold', level=1)
doc.add_paragraph('We evaluated both Sigmoid (Platt Scaling) and Isotonic Regression using CalibratedClassifierCV on a strict hold-out calibration set. By comparing their Brier Scores, the pipeline automatically selects the superior calibrator.')
doc.add_paragraph('Rather than relying on the default 0.5 decision threshold, we ran a cost-sensitive analysis across thresholds [0.1 to 0.9]. We optimized the threshold to maximize business value (penalizing costly False Negatives over False Positives), arriving at a custom operating point saved in operating_point.json.')

doc.save(r'c:\Users\777Pr\Downloads\Task12_Written_Response.docx')
print("Successfully generated Task 12 Written Response Word document!")
